from __future__ import annotations

from pathlib import Path

from src.domain.models import ArtifactKind, ArtifactRef
from src.domain.output_formatter_support import artifact_ref_with_meta
from src.domain.stata_runner import RunError


def _report_lines(*, job_id: str, artifacts: tuple[ArtifactRef, ...]) -> list[str]:
    lines = [f"job_id: {job_id}", "artifacts:"]
    lines.extend(f"- {ref.rel_path}" for ref in artifacts)
    return lines


def _write_docx_or_error(*, lines: list[str], dest_path: Path) -> RunError | None:
    try:
        from docx import Document
    except ImportError as e:
        return RunError(error_code="OUTPUT_FORMATTER_FAILED", message=str(e))
    dest_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc = Document()
        doc.add_heading("SS Output Report", level=1)
        for line in lines:
            doc.add_paragraph(line)
        doc.save(str(dest_path))
    except (OSError, ValueError) as e:
        return RunError(error_code="OUTPUT_FORMATTER_FAILED", message=str(e))
    return None


def _write_pdf_or_error(*, lines: list[str], dest_path: Path) -> RunError | None:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError as e:
        return RunError(error_code="OUTPUT_FORMATTER_FAILED", message=str(e))
    dest_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        c = canvas.Canvas(str(dest_path), pagesize=letter)
        width, height = letter
        y = height - 72
        c.setFont("Helvetica-Bold", 14)
        c.drawString(72, y, "SS Output Report")
        y -= 36
        c.setFont("Helvetica", 10)
        for line in lines:
            if y < 72:
                c.showPage()
                y = height - 72
                c.setFont("Helvetica", 10)
            c.drawString(72, y, line[:120])
            y -= 14
        c.save()
    except (OSError, ValueError) as e:
        return RunError(error_code="OUTPUT_FORMATTER_FAILED", message=str(e))
    return None


def produce_docx(
    *,
    created_at: str,
    job_id: str,
    job_dir: Path,
    formatted_dir: Path,
    artifacts: tuple[ArtifactRef, ...],
) -> tuple[ArtifactRef | None, RunError | None]:
    dest_path = formatted_dir / "report.docx"
    lines = _report_lines(job_id=job_id, artifacts=artifacts)
    err = _write_docx_or_error(lines=lines, dest_path=dest_path)
    if err is not None:
        return None, err
    return (
        artifact_ref_with_meta(
            job_dir=job_dir,
            kind=ArtifactKind.STATA_EXPORT_REPORT,
            path=dest_path,
            created_at=created_at,
            output_format="docx",
        ),
        None,
    )


def produce_pdf(
    *,
    created_at: str,
    job_id: str,
    job_dir: Path,
    formatted_dir: Path,
    artifacts: tuple[ArtifactRef, ...],
) -> tuple[ArtifactRef | None, RunError | None]:
    dest_path = formatted_dir / "report.pdf"
    lines = _report_lines(job_id=job_id, artifacts=artifacts)
    err = _write_pdf_or_error(lines=lines, dest_path=dest_path)
    if err is not None:
        return None, err
    return (
        artifact_ref_with_meta(
            job_dir=job_dir,
            kind=ArtifactKind.STATA_EXPORT_REPORT,
            path=dest_path,
            created_at=created_at,
            output_format="pdf",
        ),
        None,
    )
